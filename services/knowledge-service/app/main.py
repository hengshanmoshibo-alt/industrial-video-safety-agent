import uuid
from io import BytesIO

from fastapi import Depends, File, HTTPException, Query, UploadFile
from pydantic import BaseModel
from sqlmodel import Session, select

from aicoding_shared.db import get_session
from aicoding_shared.milvus import VectorStore
from aicoding_shared.models import ApprovalStatus, KnowledgeBase, KnowledgeChunk, KnowledgeDocument, KnowledgeVersion, User, UserRole, now_utc
from aicoding_shared.security import require_roles, tenant_from_token
from aicoding_shared.seed import ECOMMERCE_KB
from aicoding_shared.service import create_service_app
from aicoding_shared.text import chunk_text, deterministic_embedding, keywords


class DocumentCreate(BaseModel):
    title: str
    category: str = "通用"
    content: str
    source: str = "manual"
    license: str = "internal"


class OpenDatasetImport(BaseModel):
    dataset: str = "Chinese-Ambiguous-Reference"
    persist_as_draft: bool = False
    samples: int = 12


class ApprovalIn(BaseModel):
    document_id: int
    approved: bool = True
    comment: str = ""


OPEN_DATASET_CASES = {
    "Chinese-Ambiguous-Reference": {
        "source": "https://github.com/Alab-NII/Chinese-Ambiguous-Reference",
        "license": "MIT",
        "purpose": "中文指代、省略和追问测试问法扩展",
        "cases": [
            ("退款规则", "这个多久到账？"),
            ("退货换货", "如果不合适还能换吗？"),
            ("物流查询", "它现在到哪了？"),
            ("发票开具", "公司抬头怎么弄？"),
            ("修改地址", "这个地址填错了怎么办？"),
        ],
    },
    "JDDC": {
        "source": "https://github.com/SimonJYang/JDDC-Baseline-Seq2Seq",
        "license": "research-use-check-required",
        "purpose": "客服对话风格和问法扩展测试，不默认进入生产知识库",
        "cases": [
            ("订单取消", "这单我不想要了"),
            ("商品缺货", "没货了什么时候补"),
            ("支付失败", "钱扣了但是订单没成"),
            ("投诉与转人工", "我要找真人处理"),
        ],
    },
    "BQ Corpus": {
        "source": "https://www.modelscope.cn/datasets/DAMO_NLP/BQ_Corpus",
        "license": "dataset-terms-check-required",
        "purpose": "相似问识别和召回评估",
        "cases": [
            ("退款规则", "退款什么时候退回原账户"),
            ("物流查询", "快递多久能到"),
            ("优惠券", "券为什么不能用"),
        ],
    },
}


def _default_kb(session: Session, tenant_id: int) -> KnowledgeBase:
    kb = session.exec(select(KnowledgeBase).where(KnowledgeBase.tenant_id == tenant_id, KnowledgeBase.name == "电商客服知识库")).first()
    if kb is None:
        kb = KnowledgeBase(tenant_id=tenant_id, name="电商客服知识库", description="内置中文电商客服种子知识")
        session.add(kb)
        session.flush()
    return kb


def _index_document(session: Session, doc: KnowledgeDocument, explicit_keywords: list[str] | None = None) -> int:
    old = session.exec(select(KnowledgeChunk).where(KnowledgeChunk.document_id == doc.id)).all()
    store = VectorStore()
    store.delete_chunks([int(item.id) for item in old if item.id is not None])
    for item in old:
        session.delete(item)
    session.flush()
    rows: list[dict] = []
    count = 0
    for part in chunk_text(doc.content):
        vector_id = str(uuid.uuid4())
        chunk = KnowledgeChunk(
            tenant_id=doc.tenant_id,
            document_id=doc.id,
            title=doc.title,
            category=doc.category,
            content=part,
            keywords=keywords(part, explicit_keywords),
            source=doc.source,
            vector_id=vector_id,
        )
        session.add(chunk)
        session.flush()
        rows.append(
            {
                "id": int(chunk.id),
                "vector": deterministic_embedding(part),
                "tenant_id": doc.tenant_id,
                "chunk_id": chunk.id,
                "document_id": doc.id,
                "title": doc.title,
                "category": doc.category,
                "content": part,
                "source": doc.source,
            }
        )
        count += 1
    store.upsert(rows)
    return count


def _extract_upload_text(file_name: str, raw: bytes) -> str:
    suffix = file_name.rsplit(".", 1)[-1].lower() if "." in file_name else "txt"
    if suffix in {"txt", "md", "markdown"}:
        for encoding in ("utf-8", "utf-8-sig", "gb18030"):
            try:
                return raw.decode(encoding)
            except UnicodeDecodeError:
                continue
        return raw.decode("utf-8", errors="ignore")
    if suffix == "pdf":
        from pypdf import PdfReader

        reader = PdfReader(BytesIO(raw))
        return "\n".join(page.extract_text() or "" for page in reader.pages).strip()
    if suffix == "docx":
        from docx import Document

        document = Document(BytesIO(raw))
        return "\n".join(paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()).strip()
    raise HTTPException(status_code=400, detail="Unsupported file type. Use PDF, DOCX, TXT, Markdown.")


def _keyword_search(session: Session, tenant_id: int, q: str, limit: int) -> list[dict]:
    query_chars = set(q.lower())
    query = q.lower()
    synonym_groups = {
        "退款": ["退款", "退钱", "多久到账", "原路退回"],
        "退货": ["退货", "换货", "不合适"],
        "物流": ["物流", "快递", "发货", "到哪", "哪里查"],
        "发票": ["发票", "抬头", "税号"],
        "优惠券": ["优惠券", "券", "不能用"],
        "地址": ["地址", "填错", "修改地址"],
        "投诉": ["投诉", "人工", "真人", "赔偿"],
    }
    expanded_terms = {q}
    for canonical, terms in synonym_groups.items():
        if any(term in query for term in terms):
            expanded_terms.add(canonical)
            expanded_terms.update(terms)
    chunks = session.exec(select(KnowledgeChunk).where(KnowledgeChunk.tenant_id == tenant_id)).all()
    ranked = []
    for chunk in chunks:
        title = chunk.title.lower()
        category = chunk.category.lower()
        content = chunk.content.lower()
        text = f"{title} {category} {content}"
        score = len(query_chars & set(text)) / max(len(query_chars), 1)
        score += sum(1.7 for kw in chunk.keywords if kw and kw.lower() in query)
        for term in expanded_terms:
            term = term.lower()
            if not term:
                continue
            if term in title:
                score += 3.0
            if term in category:
                score += 1.8
            if term in content:
                score += 0.8
        if score > 0.12:
            ranked.append(
                (
                    score,
                    {
                        "chunk_id": chunk.id,
                        "document_id": chunk.document_id,
                        "title": chunk.title,
                        "category": chunk.category,
                        "content": chunk.content,
                        "source": chunk.source,
                        "keyword_score": round(score, 4),
                    },
                )
            )
    ranked.sort(key=lambda item: item[0], reverse=True)
    return [item for _, item in ranked[:limit]]


def _hybrid_search(session: Session, tenant_id: int, q: str, limit: int) -> dict:
    store = VectorStore()
    vector_results = store.search(deterministic_embedding(q), tenant_id=tenant_id, limit=max(limit * 2, 8))
    keyword_results = _keyword_search(session, tenant_id, q, limit=max(limit * 2, 8))

    merged: dict[int, dict] = {}
    trace = {
        "query": q,
        "embedding": "deterministic",
        "vector_store": "milvus" if store.available else "unavailable",
        "vector_candidates": len(vector_results),
        "keyword_candidates": len(keyword_results),
        "reranker": "weighted_hybrid_v1",
    }

    for index, item in enumerate(vector_results):
        chunk_id = int(item.get("chunk_id") or 0)
        if not chunk_id:
            continue
        normalized = max(float(item.get("vector_score", 0) or 0), 0.0)
        merged[chunk_id] = item | {
            "keyword_score": 0.0,
            "vector_rank": index + 1,
            "keyword_rank": None,
            "score": round(normalized * 2.2 + max(0, limit - index) * 0.03, 4),
        }

    for index, item in enumerate(keyword_results):
        chunk_id = int(item.get("chunk_id") or 0)
        if not chunk_id:
            continue
        current = merged.get(chunk_id, item | {"vector_score": 0.0, "vector_rank": None})
        keyword_score = float(item.get("keyword_score", 0) or 0)
        current.update(item)
        current["keyword_rank"] = index + 1
        current["score"] = round(float(current.get("score", 0) or 0) + keyword_score * 1.15 + max(0, limit - index) * 0.05, 4)
        merged[chunk_id] = current

    sources = sorted(merged.values(), key=lambda item: item.get("score", 0), reverse=True)[:limit]
    confidence = min(0.98, float(sources[0]["score"]) / 4.2) if sources else 0.0
    return {"sources": sources, "confidence": round(confidence, 3), "retrieval_trace": trace}


def bootstrap() -> None:
    from aicoding_shared.db import engine
    from aicoding_shared.models import Tenant

    with Session(engine) as session:
        tenant = session.exec(select(Tenant).where(Tenant.slug == "default")).first()
        if tenant is None:
            tenant = Tenant(slug="default", name="默认企业")
            session.add(tenant)
            session.flush()
        kb = _default_kb(session, tenant.id)
        for title, category, kws, content in ECOMMERCE_KB:
            exists = session.exec(select(KnowledgeDocument).where(KnowledgeDocument.tenant_id == tenant.id, KnowledgeDocument.title == title)).first()
            if exists:
                if exists.status == ApprovalStatus.published:
                    _index_document(session, exists, kws)
                continue
            doc = KnowledgeDocument(
                tenant_id=tenant.id,
                knowledge_base_id=kb.id,
                title=title,
                category=category,
                content=content,
                source="seed:ecommerce",
                license="internal_seed",
                status=ApprovalStatus.published,
                published_at=now_utc(),
            )
            session.add(doc)
            session.flush()
            session.add(KnowledgeVersion(tenant_id=tenant.id, document_id=doc.id, version=1, content=content, status=ApprovalStatus.published))
            _index_document(session, doc, kws)
        session.commit()


app = create_service_app("knowledge-service", bootstrap=bootstrap)


@app.get("/kb/documents")
def list_documents(tenant=Depends(tenant_from_token), session: Session = Depends(get_session)):
    return session.exec(select(KnowledgeDocument).where(KnowledgeDocument.tenant_id == tenant.id).order_by(KnowledgeDocument.id.desc())).all()


@app.post("/kb/documents")
def create_document(payload: DocumentCreate, user: User = Depends(require_roles([UserRole.admin, UserRole.kb_manager])), session: Session = Depends(get_session)):
    kb = _default_kb(session, user.tenant_id)
    doc = KnowledgeDocument(
        tenant_id=user.tenant_id,
        knowledge_base_id=kb.id,
        title=payload.title,
        category=payload.category,
        content=payload.content,
        source=payload.source,
        license=payload.license,
        status=ApprovalStatus.draft,
    )
    session.add(doc)
    session.flush()
    session.add(KnowledgeVersion(tenant_id=user.tenant_id, document_id=doc.id, version=1, content=doc.content, status=ApprovalStatus.draft))
    session.commit()
    session.refresh(doc)
    return doc


@app.post("/kb/documents/upload")
async def upload_document(
    title: str,
    category: str = "通用",
    source: str = "upload",
    license: str = "internal",
    file: UploadFile = File(...),
    user: User = Depends(require_roles([UserRole.admin, UserRole.kb_manager])),
    session: Session = Depends(get_session),
):
    raw = await file.read()
    content = _extract_upload_text(file.filename or title, raw)
    if not content.strip():
        raise HTTPException(status_code=400, detail="Uploaded document has no extractable text")
    kb = _default_kb(session, user.tenant_id)
    doc = KnowledgeDocument(
        tenant_id=user.tenant_id,
        knowledge_base_id=kb.id,
        title=title,
        category=category,
        content=content,
        source=source,
        license=license,
        status=ApprovalStatus.draft,
    )
    session.add(doc)
    session.flush()
    session.add(KnowledgeVersion(tenant_id=user.tenant_id, document_id=doc.id, version=1, content=content, status=ApprovalStatus.draft))
    session.commit()
    session.refresh(doc)
    return {"document": doc, "characters": len(content), "filename": file.filename}


@app.post("/kb/documents/{document_id}/reindex")
def reindex(document_id: int, user: User = Depends(require_roles([UserRole.admin, UserRole.kb_manager])), session: Session = Depends(get_session)):
    doc = session.get(KnowledgeDocument, document_id)
    if doc is None or doc.tenant_id != user.tenant_id:
        raise HTTPException(status_code=404, detail="Document not found")
    chunks = _index_document(session, doc)
    session.commit()
    return {"document_id": document_id, "chunks": chunks}


@app.delete("/kb/documents/{document_id}")
def delete_document(document_id: int, user: User = Depends(require_roles([UserRole.admin, UserRole.kb_manager])), session: Session = Depends(get_session)):
    doc = session.get(KnowledgeDocument, document_id)
    if doc is None or doc.tenant_id != user.tenant_id:
        raise HTTPException(status_code=404, detail="Document not found")
    chunks = session.exec(select(KnowledgeChunk).where(KnowledgeChunk.tenant_id == user.tenant_id, KnowledgeChunk.document_id == document_id)).all()
    VectorStore().delete_chunks([int(item.id) for item in chunks if item.id is not None])
    for chunk in chunks:
        session.delete(chunk)
    doc.is_active = False
    doc.status = ApprovalStatus.rejected
    doc.updated_at = now_utc()
    session.add(doc)
    session.commit()
    return {"document_id": document_id, "deleted": True, "chunks_removed": len(chunks)}


@app.get("/kb/versions")
def list_versions(tenant=Depends(tenant_from_token), session: Session = Depends(get_session)):
    return session.exec(select(KnowledgeVersion).where(KnowledgeVersion.tenant_id == tenant.id).order_by(KnowledgeVersion.id.desc())).all()


@app.get("/kb/approvals")
def list_approvals(user: User = Depends(require_roles([UserRole.admin, UserRole.kb_manager])), session: Session = Depends(get_session)):
    return session.exec(select(KnowledgeDocument).where(KnowledgeDocument.tenant_id == user.tenant_id, KnowledgeDocument.status == ApprovalStatus.pending)).all()


@app.post("/kb/approvals")
def approve(payload: ApprovalIn, user: User = Depends(require_roles([UserRole.admin, UserRole.kb_manager])), session: Session = Depends(get_session)):
    doc = session.get(KnowledgeDocument, payload.document_id)
    if doc is None or doc.tenant_id != user.tenant_id:
        raise HTTPException(status_code=404, detail="Document not found")
    doc.status = ApprovalStatus.approved if payload.approved else ApprovalStatus.rejected
    session.add(doc)
    session.add(KnowledgeVersion(tenant_id=user.tenant_id, document_id=doc.id, version=doc.version, content=doc.content, status=doc.status, reviewer_id=user.id))
    session.commit()
    session.refresh(doc)
    return doc


@app.post("/kb/publish")
def publish(payload: ApprovalIn, user: User = Depends(require_roles([UserRole.admin, UserRole.kb_manager])), session: Session = Depends(get_session)):
    doc = session.get(KnowledgeDocument, payload.document_id)
    if doc is None or doc.tenant_id != user.tenant_id:
        raise HTTPException(status_code=404, detail="Document not found")
    doc.version += 1
    doc.status = ApprovalStatus.published
    doc.published_at = now_utc()
    _index_document(session, doc)
    session.add(doc)
    session.add(KnowledgeVersion(tenant_id=user.tenant_id, document_id=doc.id, version=doc.version, content=doc.content, status=ApprovalStatus.published, reviewer_id=user.id))
    session.commit()
    session.refresh(doc)
    return doc


@app.get("/kb/search")
def search(q: str, tenant=Depends(tenant_from_token), session: Session = Depends(get_session), limit: int = 5, include_trace: bool = Query(default=False)):
    result = _hybrid_search(session, tenant.id, q, limit)
    if include_trace:
        return result
    return result["sources"]


@app.post("/kb/seed/ecommerce")
def seed(user: User = Depends(require_roles([UserRole.admin, UserRole.kb_manager])), session: Session = Depends(get_session)):
    kb = _default_kb(session, user.tenant_id)
    created = 0
    reindexed = 0
    for title, category, kws, content in ECOMMERCE_KB:
        doc = session.exec(select(KnowledgeDocument).where(KnowledgeDocument.tenant_id == user.tenant_id, KnowledgeDocument.title == title)).first()
        if doc is None:
            doc = KnowledgeDocument(
                tenant_id=user.tenant_id,
                knowledge_base_id=kb.id,
                title=title,
                category=category,
                content=content,
                source="seed:ecommerce",
                license="internal_seed",
                status=ApprovalStatus.published,
                published_at=now_utc(),
            )
            session.add(doc)
            session.flush()
            session.add(KnowledgeVersion(tenant_id=user.tenant_id, document_id=doc.id, version=1, content=content, status=ApprovalStatus.published))
            created += 1
        _index_document(session, doc, kws)
        reindexed += 1
    session.commit()
    return {"documents_created": created, "documents_reindexed": reindexed}


@app.post("/kb/import/open-dataset")
def import_open_dataset(payload: OpenDatasetImport, user: User = Depends(require_roles([UserRole.admin, UserRole.kb_manager])), session: Session = Depends(get_session)):
    dataset = OPEN_DATASET_CASES.get(payload.dataset)
    if dataset is None:
        raise HTTPException(status_code=404, detail="Unsupported dataset")
    cases = [
        {"label": label, "question": question, "expected_use": "test_or_eval_only"}
        for label, question in dataset["cases"][: max(1, min(payload.samples, 50))]
    ]
    created = 0
    if payload.persist_as_draft:
        kb = _default_kb(session, user.tenant_id)
        for case in cases:
            doc = KnowledgeDocument(
                tenant_id=user.tenant_id,
                knowledge_base_id=kb.id,
                title=f"测试问法：{case['label']}",
                category="测试评估",
                content=f"问题：{case['question']}\n期望命中：{case['label']}",
                source=f"open-dataset:{payload.dataset}",
                license=str(dataset["license"]),
                status=ApprovalStatus.draft,
                is_active=False,
            )
            session.add(doc)
            session.flush()
            session.add(KnowledgeVersion(tenant_id=user.tenant_id, document_id=doc.id, version=1, content=doc.content, status=ApprovalStatus.draft))
            created += 1
        session.commit()
    return {
        "dataset": payload.dataset,
        "source": dataset["source"],
        "license": dataset["license"],
        "purpose": dataset["purpose"],
        "persisted_as_draft": payload.persist_as_draft,
        "documents_created": created,
        "cases": cases,
    }
